#!/usr/bin/python
# -*- coding: utf-8 -*-
import string
import time
import Levenshtein
import pandas as pd
from docx import Document
from collections import defaultdict
import re

# Characters to be removed from the words during text preprocessing
punctuation_to_remove = '"!.,“”;[]():﻿«»…#*―_<>%...  =—§ʿ–'

# Common word parts that should be considered as part of the root word
# and not treated as separate words
appendices = ['dan', 'den', 'dir', 'dir', 'inin', 'nde', 'nden', 'nin', 'nun', 'nın', 'tan', 'ten', 'daki']

# check if a word contains any digit
def has_digit(word):
    return any(char.isdigit() for char in word)

# handle apostrophes at the beginning or end of a word
def apostrophe(word):
    # Handle special cases where the apostrophe is at the beginning or end of the word
    if len(word) > 0 and word[0] in {'ʻ', '‘', '’', '–', '?', 'ʿ', '–', ' '}:
        word = word[1:]
        return apostrophe(word)
    if len(word) > 0 and word[-1] in {'ʻ', '‘', '’', '?', ' ', '–'}:
        word = word[0:-1]
        return apostrophe(word)
    return word

# standartise Unicode characters and handle Turkish characters
def preprocess_line(line):
    line = line.lower().replace("-", " ")
    line = line.lower().replace(".", " ")
    line = line.lower().replace("?", " ")
    line = re.sub(r'[\u00AD\u002D\u2011]+', '', line)
    line = line.replace('\u0075\u0308', 'ü')
    line = line.replace('\u0069\u0307' ,'i')
    line = line.replace('\u0067\u0306' ,'ğ') 
    line = line.replace('\u0063\u0327' ,'ç') 
    line = line.replace('\u0073\u0327' ,'ş')
    return line

# Set to store exception words found during text preprocessing
exceptions = set()

# Text preprocessing to remove unnecessary punctuation and extract valid words
def preprocess_text(text):
    words = set()
    for line in text.splitlines():
        line = line.strip()
        if line:
            line = preprocess_line(line)
            words_in_line = line.lower().split()
            for word in words_in_line:
                if len(word) > 0 and word and word != '\n':
                    word = word.translate(str.maketrans('', '', punctuation_to_remove))
                    word = word.split("'")[0]
                    word = apostrophe(word)
                    if not has_digit(word) and len(word) > 3 and word != "discard" and word != "disccard" and word not in appendices and '?' not in word:
                        #word = str(len(word)) + " -> " + word 
                        words.add(word)
                    elif '?' in word and not has_digit(word) and len(word) > 3 and word != "discard" and word != "disccard" and word not in appendices:
                        exceptions.add(word)

    return sorted(words, key=len)

# find the top n most similar words to a given word
# based on Levenshtein distance with a maximum allowed difference
def find_most_similar_words(word, unique_word_list, n=3, max_difference=1):
    similarities = {}
    for unique_word in unique_word_list:
        # Avoid self-comparison and make an early exit if necessary
        if word != unique_word and abs(len(word) - len(unique_word)) <= max_difference:
            distance = Levenshtein.distance(word, unique_word)
            if distance > max_difference:
                continue  # Skip this word if the difference is too large
            similarity = 1 - distance / max(len(word), len(unique_word))
            similarities[unique_word] = similarity

    # Get the top n most similar words
    sorted_words = sorted(similarities.items(), key=lambda x: x[1], reverse=True)

    # Check if there are enough similar words to return
    num_similar_words = len(sorted_words)
    if num_similar_words < n:
        n = num_similar_words

    return [word for word, _ in sorted_words[:n]]

# check if a paragraph in the docx file is bold
def is_bold(paragraph):
    for run in paragraph.runs:
        if run.bold:
            return True
    return False

unique_words = set()
start_time = time.time()

# Read the docx file and extract unique words
doc = Document('unique_words.docx') # take filename as main argument later
#doc = Document('Sample_Doc.docx')

for paragraph in doc.paragraphs:
    if not is_bold(paragraph):
        text = paragraph.text
        if text:
            unique_words.update(preprocess_text(text))

# Save the unique words
unique_words_list = sorted(list(unique_words))
df_unique = pd.Series(unique_words_list)
output_file = 'result_of_uniquewords.txt'
with open(output_file, 'w', encoding='utf-8') as file:
    file.write(df_unique.to_string())

# Save the exception words
exceptions_list = sorted(list(exceptions))
df_exception = pd.Series(exceptions_list)
output_file = 'result_of_exceptions.txt'
with open(output_file, 'w', encoding='utf-8') as file:
    file.write(df_exception.to_string())


print("Code is running to find the similar words for ", len(unique_words) , " unique words.")

# Find similar words for each unique word and save to a DataFrame
similar_words_dict = defaultdict(list)
for word in unique_words:
    similar_words = find_most_similar_words(word, unique_words)
    similar_words_dict[word] = similar_words

df = pd.DataFrame(similar_words_dict.items(), columns=['Word', 'Similar Words'])
df = df.sort_values(by='Word')

output_file = 'similar_unique_words.txt'
with open(output_file, 'w', encoding='utf-8') as file:
    file.write(df.to_string(index=False))

end_time = time.time()
elapsed_time = end_time - start_time

# Print Benchmark
print("Process is completed.")
print(f"Results saved to {output_file}")

hours = int(elapsed_time // 3600)
minutes = int((elapsed_time % 3600) // 60)
seconds = elapsed_time % 60

print(f"Elapsed time: {hours} hours, {minutes} minutes, {seconds:.2f} seconds")